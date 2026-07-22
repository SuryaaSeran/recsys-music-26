"""Render paper_wikis/paper_draft.md into the ACM acmart Word template.

Loads the interim ACM template (for its named styles + page setup), clears the
sample body, and rebuilds the paper using acmart styles (Title_document, AbsHead,
Abstract, CCSHead, KeyWordHead, Head1/2/3, Para, AckHead, ReferenceHead,
Bib_entry, TableCaption, Table Grid). The template's body section is two-column
(~3.33in per column); wide data tables are wrapped in continuous section breaks
so they span the full page width instead of overflowing a single column.
"""
import copy
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[2]
TEMPLATE = Path.home() / ".claude/uploads/95ce8712-e6ed-46cb-814b-7ef94e493b61/56402faa-interimlayout.docx"
SRC = ROOT / "paper_wikis/paper_draft.md"
OUT = ROOT / "paper_wikis/paper.docx"

USABLE_WIDTH_IN = 7.0  # page width 8.5in - 0.75in margins each side
COL_GUTTER_TWIPS = 480

ALIGN = {"l": WD_ALIGN_PARAGRAPH.LEFT, "c": WD_ALIGN_PARAGRAPH.CENTER, "r": WD_ALIGN_PARAGRAPH.RIGHT}

# per-table column widths (inches, sum to USABLE_WIDTH_IN) and per-column alignment,
# keyed by a lowercase substring of the table caption
TABLE_SPECS = {
    "nine recall sources": {"widths": [0.5, 1.6, 1.3, 3.6], "align": ["c", "l", "c", "l"]},
    "rescorer feature groups": {"widths": [1.5, 5.5], "align": ["l", "l"]},
    "development ndcg@20 after each load-bearing change": {"widths": [5.4, 1.6], "align": ["l", "r"]},
    "reranking attempts": {"widths": [2.0, 2.8, 2.2], "align": ["l", "l", "l"]},
    "dev ndcg@20 by goal category": {"widths": [0.6, 4.9, 1.5], "align": ["c", "l", "r"]},
    "measured latency and footprint": {"widths": [5.0, 2.0], "align": ["l", "r"]},
}


def table_spec(caption, ncols):
    cap = (caption or "").lower()
    for key, spec in TABLE_SPECS.items():
        if key in cap and len(spec["widths"]) == ncols:
            return spec
    even = USABLE_WIDTH_IN / ncols
    return {"widths": [even] * ncols, "align": ["l"] * ncols}


def set_table_grid(tbl, widths_in):
    """Set the table-level column grid (outside tcPr, no ordering constraints)."""
    tblPr = tbl._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    grid = tbl._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for gc, w in zip(grid.findall(qn("w:gridCol")), widths_in):
            gc.set(qn("w:w"), str(int(w * 1440)))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER


def set_cell_border(cell, **edges):
    """edges: top/bottom -> dict(sz=eighths-of-a-point, val='single', color='hex')."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge, spec in edges.items():
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), spec.get("val", "single"))
        el.set(qn("w:sz"), str(spec.get("sz", 4)))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), spec.get("color", "000000"))
        borders.append(el)
    tcPr.append(borders)


def shade_cell(cell, hexfill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexfill)
    tcPr.append(shd)


def repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    trPr.append(header)


def set_cell_margins(tbl, top=40, bottom=40, left=100, right=100):
    tblPr = tbl._tbl.tblPr
    mar = OxmlElement("w:tblCellMar")
    for edge, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tblPr.append(mar)


def insert_section_break(doc, cols):
    """End the current section here with the given column count (continuous
    break) so content before/after this point can use a different column
    count than content that follows. Copies page size/margins from the
    document's trailing sectPr so page geometry stays consistent."""
    body_sectPr = doc.sections[-1]._sectPr
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    sectPr = OxmlElement("w:sectPr")

    typ = OxmlElement("w:type")
    typ.set(qn("w:val"), "continuous")
    sectPr.append(typ)

    for tag in ("w:pgSz", "w:pgMar"):
        el = body_sectPr.find(qn(tag))
        if el is not None:
            sectPr.append(copy.deepcopy(el))

    colsEl = OxmlElement("w:cols")
    colsEl.set(qn("w:num"), str(cols))
    if cols > 1:
        colsEl.set(qn("w:space"), str(COL_GUTTER_TWIPS))
    sectPr.append(colsEl)

    pPr.append(sectPr)
    return p


def has_style(doc, name):
    try:
        _ = doc.styles[name]
        return True
    except KeyError:
        return False


def add_runs(paragraph, text):
    """Parse inline **bold**, *italic*, `mono` and add formatted runs."""
    token = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")
    for part in token.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2]); r.bold = True
        elif part.startswith("*") and part.endswith("*"):
            r = paragraph.add_run(part[1:-1]); r.italic = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1]); r.font.name = "Courier New"
        else:
            paragraph.add_run(part)


def main():
    doc = Document(str(TEMPLATE))

    # style resolution with fallbacks
    S = {
        "title": "Title_document" if has_style(doc, "Title_document") else "Title",
        "authors": "Authors",
        "affil": "Affiliation" if has_style(doc, "Affiliation") else "Authors",
        "abshead": "AbsHead", "abstract": "Abstract",
        "ccshead": "CCSHead", "ccs": "CCSDescription",
        "kwhead": "KeyWordHead", "kw": "KeyWords",
        "h1": "Head1", "h2": "Head2", "h3": "Head3",
        "para": "Para", "list": "List Paragraph",
        "quote": "Quote" if has_style(doc, "Quote") else "Para",
        "code": "programCode_display" if has_style(doc, "programCode_display") else "Para",
        "ackhead": "AckHead", "ackpara": "AckPara",
        "refhead": "ReferenceHead", "bib": "Bib_entry" if has_style(doc, "Bib_entry") else "Para",
        "tcap": "TableCaption" if has_style(doc, "TableCaption") else "Caption",
    }

    # clear sample body paragraphs and tables
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)
    for t in list(doc.tables):
        t._element.getparent().remove(t._element)

    def add(style, text=""):
        p = doc.add_paragraph(style=S[style] if style in S else style)
        if text:
            add_runs(p, text)
        return p

    def add_table(rows, caption=None):
        header, body = rows[0], rows[1:]
        ncols = len(header)
        spec = table_spec(caption, ncols)
        widths, aligns = spec["widths"], spec["align"]

        tbl = doc.add_table(rows=1, cols=ncols)
        tbl.style = "Normal Table" if has_style(doc, "Normal Table") else "Table Grid"
        set_cell_margins(tbl)

        HEADER_RULE = {"sz": 12, "val": "single", "color": "000000"}
        MID_RULE = {"sz": 4, "val": "single", "color": "000000"}

        for j, cell in enumerate(header):
            c = tbl.rows[0].cells[j]
            c.width = Inches(widths[j])
            c.paragraphs[0].text = ""
            c.paragraphs[0].alignment = ALIGN.get(aligns[j], WD_ALIGN_PARAGRAPH.LEFT)
            r = c.paragraphs[0].add_run(cell); r.bold = True
            r.font.size = Pt(9)
            shade_cell(c, "E7E7E7")
            set_cell_border(c, top=HEADER_RULE, bottom=MID_RULE)
        repeat_header(tbl.rows[0])

        for ridx, row in enumerate(body):
            is_last = ridx == len(body) - 1
            cells = tbl.add_row().cells
            for j, cell in enumerate(row):
                if j >= len(cells):
                    break
                c = cells[j]
                c.width = Inches(widths[j])
                c.paragraphs[0].text = ""
                c.paragraphs[0].alignment = ALIGN.get(aligns[j], WD_ALIGN_PARAGRAPH.LEFT)
                add_runs(c.paragraphs[0], cell)
                for run in c.paragraphs[0].runs:
                    run.font.size = Pt(9)
                if is_last:
                    set_cell_border(c, bottom=HEADER_RULE)

        set_table_grid(tbl, widths)

    lines = SRC.read_text().splitlines()
    i = 0
    seen_title = False
    section = "front"  # front|abstract|ccs|kw|body|ack|ref
    para_buf = []
    last_caption = None

    def flush_para():
        nonlocal para_buf
        if not para_buf:
            return
        text = " ".join(para_buf).strip()
        para_buf = []
        if not text:
            return
        if section == "abstract":
            add("abstract", text)
        elif section == "ccs":
            add("ccs", text)
        elif section == "kw":
            add("kw", text)
        elif section == "ack":
            add("ackpara", text)
        elif section == "ref":
            add("bib", text)
        else:
            add("para", text)

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # skip the instructional italic note (may span multiple lines)
        if stripped.startswith("*("):
            while i < len(lines) and not lines[i].strip().endswith(")*"):
                i += 1
            i += 1  # consume the closing line
            continue
        if stripped == "---":
            flush_para(); i += 1; continue

        # table caption: **Table N: ...**  -> remembered, added inside the
        # table's own full-width section (see table block below). May wrap
        # across multiple source lines before the closing "**".
        if re.match(r"^\*\*Table \d+[a-z]?:", stripped):
            flush_para()
            parts = [stripped]
            while not parts[-1].endswith("**") and i + 1 < len(lines):
                i += 1
                parts.append(lines[i].strip())
            caption_text = " ".join(parts)
            caption_text = re.sub(r"^\*\*", "", caption_text)
            caption_text = re.sub(r"\*\*$", "", caption_text)
            last_caption = caption_text.strip()
            i += 1
            continue

        # fenced code block
        if stripped.startswith("```"):
            flush_para()
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                p = doc.add_paragraph(style=S["code"])
                r = p.add_run(lines[i]); r.font.name = "Courier New"; r.font.size = Pt(9)
                i += 1
            i += 1
            continue

        # table block
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s:|-]+\|$", lines[i + 1].strip()):
            flush_para()
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not re.match(r"^[\s:|-]+$", "".join(cells)):
                    rows.append(cells)
                i += 1
            insert_section_break(doc, cols=1)
            if last_caption:
                add("tcap", last_caption)
            add_table(rows, caption=last_caption)
            last_caption = None
            insert_section_break(doc, cols=2)
            continue

        # headings
        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            flush_para()
            level, htext = len(m.group(1)), m.group(2).strip()
            if level == 1 and not seen_title:
                add("title", htext); seen_title = True
                section = "front"
            elif htext.upper() == "ABSTRACT":
                add("abshead", "ABSTRACT"); section = "abstract"
            elif htext.upper().startswith("CCS"):
                add("ccshead", "CCS CONCEPTS"); section = "ccs"
            elif htext.upper().startswith("KEYWORD"):
                add("kwhead", "KEYWORDS"); section = "kw"
            elif htext.upper().startswith("ACKNOWLE"):
                add("ackhead", "ACKNOWLEDGMENTS"); section = "ack"
            elif htext.upper().startswith("REFERENCE"):
                add("refhead", "REFERENCES"); section = "ref"
            else:
                section = "body"
                add({2: "h1", 3: "h2", 4: "h3"}.get(level, "h1"), htext)
            i += 1
            continue

        # front-matter author lines
        if section == "front" and stripped.startswith("**Authors:**"):
            flush_para()
            add("authors", stripped.replace("**Authors:**", "").strip())
            i += 1; continue
        if section == "front" and stripped.startswith("Department Name"):
            flush_para()
            add("affil", stripped)
            i += 1; continue

        # blockquote (render literally; the composite formula contains '*')
        if stripped.startswith(">"):
            flush_para()
            p = doc.add_paragraph(style=S["quote"])
            p.add_run(stripped.lstrip("> ").strip())
            i += 1; continue

        # list items (gather indented continuation lines into one item)
        lm = re.match(r"^(\d+\.|[-*])\s+(.*)$", stripped)
        if lm:
            flush_para()
            item = lm.group(2)
            j = i + 1
            while (j < len(lines) and lines[j].strip()
                   and lines[j][:1] == " "
                   and not re.match(r"^\s*(\d+\.|[-*])\s+", lines[j])
                   and lines[j].strip()[0] not in "#|`"):
                item += " " + lines[j].strip()
                j += 1
            style = "ccs" if section == "ccs" else ("kw" if section == "kw" else "list")
            p = doc.add_paragraph(style=S[style])
            add_runs(p, item)
            i = j
            continue

        # blank line ends a paragraph
        if not stripped:
            flush_para(); i += 1; continue

        para_buf.append(stripped)
        i += 1

    flush_para()
    doc.save(str(OUT))
    print("wrote", OUT)
    # sanity
    d2 = Document(str(OUT))
    print("paragraphs:", len(d2.paragraphs), "tables:", len(d2.tables))


if __name__ == "__main__":
    main()
